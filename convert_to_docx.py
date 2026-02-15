"""
Convert PROJECT_DOCUMENTATION.md to .docx format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines (but add spacing)
        if not line.strip():
            if i > 0 and not lines[i-1].startswith('#'):
                doc.add_paragraph()
            i += 1
            continue
        
        # Skip horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_page_break()
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Main title - H1
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('## '):
            # Heading 2
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True
            p.space_before = Pt(12)
            p.space_after = Pt(6)
            
        elif line.startswith('### '):
            # Heading 3
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
            
        elif line.startswith('#### '):
            # Heading 4
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
            
        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.5)
            
        # Handle tables (simple detection)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Start of table
            table_lines = []
            while i < len(lines) and ('|' in lines[i] or lines[i].strip().startswith('|-')):
                if not lines[i].strip().startswith('|-'):  # Skip separator line
                    table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # Parse table
                rows = []
                for tline in table_lines:
                    cells = [cell.strip() for cell in tline.split('|') if cell.strip()]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Bold header row
                            if row_idx == 0:
                                cell.paragraphs[0].runs[0].font.bold = True
                    
                    doc.add_paragraph()  # Space after table
            continue
            
        # Handle regular paragraphs
        else:
            text = line.strip()
            if text:
                # Clean markdown formatting
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic  
                text = re.sub(r'`(.*?)`', r'\1', text)  # Code
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
                
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Document converted successfully: {docx_file}")

if __name__ == "__main__":
    md_file = "PROJECT_DOCUMENTATION.md"
    docx_file = "InnStay_Project_Documentation.docx"
    
    try:
        parse_markdown_to_docx(md_file, docx_file)
        print(f"\n✓ Conversion complete!")
        print(f"  Input:  {md_file}")
        print(f"  Output: {docx_file}")
    except Exception as e:
        print(f"✗ Error: {e}")
